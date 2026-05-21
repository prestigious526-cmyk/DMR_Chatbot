import os
import re
from docx import Document

def convert_docx_text_and_tables_to_md(docx_path, output_md_path):
    # Load the Word document
    doc = Document(docx_path)
    md_content = []

    print(f"Parsing '{docx_path}' sequentially for text and tables...")

    # Iterate through structural elements in order of appearance
    for element in doc.element.body:
        
        # 1. HANDLE PARAGRAPHS (Headings, Bullet Lists, Plain Text)
        if element.tag.endswith('p'):
            p_objects = [p for p in doc.paragraphs if p._element == element]
            if not p_objects:
                continue
            paragraph = p_objects[0]
            text = paragraph.text.strip()
            
            # Skip completely empty paragraphs
            if not text:
                continue
            
            # Map default Word Heading styles to Markdown headers
            if paragraph.style.name.startswith('Heading 1'):
                md_content.append(f"\n# {text}\n")
            elif paragraph.style.name.startswith('Heading 2'):
                md_content.append(f"\n## {text}\n")
            elif paragraph.style.name.startswith('Heading 3'):
                md_content.append(f"\n### {text}\n")
            else:
                # Format list items uniformly
                if paragraph.style.name.startswith('List Bullet') or text.startswith('•'):
                    cleaned_bullet = text.lstrip('• ').strip()
                    md_content.append(f"* {cleaned_bullet}")
                else:
                    # Append standard paragraph body text
                    md_content.append(f"\n{text}\n")

        # 2. HANDLE TABLES (Convert structurally to standard Markdown grid strings)
        elif element.tag.endswith('tbl'):
            t_objects = [t for t in doc.tables if t._element == element]
            if not t_objects:
                continue
            table = t_objects[0]
            
            table_md = []
            for i, row in enumerate(table.rows):
                # Clean nested newlines within cells to keep markdown row integrity
                row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                
                # Render pipeline-delimited row
                table_md.append("| " + " | ".join(row_data) + " |")
                
                # Append the essential Markdown structural divider right under the header
                if i == 0:
                    dividers = ["---"] * len(row_data)
                    table_md.append("| " + " | ".join(dividers) + " |")
            
            md_content.append("\n" + "\n".join(table_md) + "\n")

    # Combine text array and clean structural formatting gaps
    full_markdown = "\n".join(md_content)
    
    # Compress excessive consecutive newlines to maintain strict chunk density
    full_markdown = re.sub(r'\n{3,}', '\n\n', full_markdown)

    # Output to the final .md target file
    with open(output_md_path, "w", encoding="utf-8") as f:
        f.write(full_markdown.strip() + "\n")
        
    print(f"Successfully converted! Generated text-only Markdown document at: '{output_md_path}'")

# --- Direct Script Execution ---
if __name__ == "__main__":
    # Ensure input path matches your file tree location
    input_file = "PART-I,Vol-I.docx"
    output_file = "Secure_Part1.md"
    
    convert_docx_text_and_tables_to_md(input_file, output_file)